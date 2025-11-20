from typing import List, Dict, Any, Optional
import io
from pathlib import Path
from loguru import logger
import mimetypes


class DocumentService:
    """Service for parsing and extracting content from various document formats"""

    def __init__(self):
        # Map of supported file extensions to parsing methods
        self.supported_formats = {
            '.txt': self._parse_text,
            '.md': self._parse_text,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_doc,
            '.json': self._parse_json,
            '.csv': self._parse_csv,
            '.xlsx': self._parse_excel,
            '.xls': self._parse_excel,
        }

    def is_supported(self, filename: str) -> bool:
        """Check if the file format is supported"""
        ext = Path(filename).suffix.lower()
        return ext in self.supported_formats

    async def parse_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse document and extract text content

        Args:
            file_content: Binary content of the file
            filename: Name of the file (used to determine format)

        Returns:
            Dictionary with extracted content and metadata
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}")

        try:
            parser = self.supported_formats[ext]
            content = await parser(file_content, filename)

            return {
                "filename": filename,
                "format": ext,
                "content": content,
                "length": len(content),
                "mime_type": mimetypes.guess_type(filename)[0] or "application/octet-stream"
            }
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {str(e)}")
            raise ValueError(f"Failed to parse document: {str(e)}")

    async def _parse_text(self, file_content: bytes, filename: str) -> str:
        """Parse plain text or markdown files"""
        try:
            # Try UTF-8 first, then fall back to other encodings
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1')
        except Exception as e:
            raise ValueError(f"Failed to decode text file: {str(e)}")

    async def _parse_pdf(self, file_content: bytes, filename: str) -> str:
        """Parse PDF files"""
        try:
            import pypdf

            pdf_file = io.BytesIO(file_content)
            reader = pypdf.PdfReader(pdf_file)

            text_content = []
            for page in reader.pages:
                text_content.append(page.extract_text())

            return "\n\n".join(text_content)
        except ImportError:
            raise ValueError("PDF parsing requires 'pypdf' package. Please install it.")
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    async def _parse_docx(self, file_content: bytes, filename: str) -> str:
        """Parse DOCX files"""
        try:
            import docx

            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)

            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            return "\n\n".join(text_content)
        except ImportError:
            raise ValueError("DOCX parsing requires 'python-docx' package. Please install it.")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    async def _parse_doc(self, file_content: bytes, filename: str) -> str:
        """Parse legacy DOC files"""
        try:
            import textract

            # textract can handle .doc files
            text = textract.process(io.BytesIO(file_content))
            return text.decode('utf-8')
        except ImportError:
            raise ValueError("DOC parsing requires 'textract' package. Please install it.")
        except Exception as e:
            raise ValueError(f"Failed to parse DOC: {str(e)}")

    async def _parse_json(self, file_content: bytes, filename: str) -> str:
        """Parse JSON files"""
        try:
            import json

            data = json.loads(file_content.decode('utf-8'))
            # Convert JSON to readable text format
            return json.dumps(data, indent=2)
        except Exception as e:
            raise ValueError(f"Failed to parse JSON: {str(e)}")

    async def _parse_csv(self, file_content: bytes, filename: str) -> str:
        """Parse CSV files"""
        try:
            import csv

            text_content = file_content.decode('utf-8')
            csv_file = io.StringIO(text_content)
            reader = csv.reader(csv_file)

            rows = []
            for row in reader:
                rows.append(" | ".join(row))

            return "\n".join(rows)
        except Exception as e:
            raise ValueError(f"Failed to parse CSV: {str(e)}")

    async def _parse_excel(self, file_content: bytes, filename: str) -> str:
        """Parse Excel files (XLSX/XLS)"""
        try:
            import openpyxl

            excel_file = io.BytesIO(file_content)
            workbook = openpyxl.load_workbook(excel_file, data_only=True)

            text_content = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    if row_text.strip():
                        text_content.append(row_text)

                text_content.append("")  # Empty line between sheets

            return "\n".join(text_content)
        except ImportError:
            raise ValueError("Excel parsing requires 'openpyxl' package. Please install it.")
        except Exception as e:
            raise ValueError(f"Failed to parse Excel: {str(e)}")

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())


# Singleton instance
document_service = DocumentService()
